from apps.imports.parsers import DocumentParser


class DocxParser(DocumentParser):
    extension = "docx"

    def parse(self, file_path: str) -> str:
        try:
            from docx import Document
        except ImportError:
            raise RuntimeError(
                "python-docx is required for DOCX import. "
                "Install it: pip install python-docx"
            )

        doc = Document(file_path)
        parts: list[str] = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)

        # Also extract table cell text
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)

        return "\n".join(parts)
