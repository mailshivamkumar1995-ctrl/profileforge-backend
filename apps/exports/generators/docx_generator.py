from apps.exports.generators import DocumentGenerator


class DocxGenerator(DocumentGenerator):
    """
    Profile data → DOCX generator using python-docx.

    Supports two layouts:
    - "classic": Two-column Georgia serif layout matching the Classic HTML template.
    - default (ATS Clean): Single-column clean professional layout.

    DOCX cannot render CSS, so each layout is built programmatically from
    profile_data to match the visual structure of the corresponding HTML template.
    """

    format = "docx"

    def generate(self, html_content: str, profile_data: dict | None = None, template_slug: str | None = None) -> bytes:
        if not profile_data:
            raise ValueError("profile_data is required for DOCX generation.")
        try:
            from docx import Document
            from docx.shared import Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            import io
            import re
        except ImportError:
            raise RuntimeError(
                "python-docx is required for DOCX export. "
                "Install it: pip install python-docx"
            )

        doc = Document()

        # ── Classic two-column layout ─────────────────────────────────────────
        if template_slug == "classic":
            NAVY = RGBColor(0x2C, 0x3E, 0x50)
            DARK = RGBColor(0x1A, 0x1A, 0x1A)
            GRAY = RGBColor(0x55, 0x55, 0x55)
            MID = RGBColor(0x44, 0x44, 0x44)

            from docx.shared import Inches
            from docx.enum.text import WD_LINE_SPACING

            def _compact(p, before=0, after=0):
                p.paragraph_format.space_before = Pt(before)
                p.paragraph_format.space_after = Pt(after)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                p.paragraph_format.line_spacing = 1.15

            for section in doc.sections:
                section.page_width = Inches(8.5)
                section.page_height = Inches(11)
                section.top_margin = Cm(0.76)   # 0.3in — matches @page margin in classic.html
                section.bottom_margin = Cm(0.76)
                section.left_margin = Cm(1.016)  # 0.4in
                section.right_margin = Cm(1.016)

            def _section_title(cell, title):
                p = cell.add_paragraph()
                _compact(p, before=4, after=1)
                run = p.add_run(title.upper())
                run.font.name = "Georgia"
                run.font.size = Pt(8)
                run.font.bold = True
                run.font.color.rgb = NAVY
                pPr = p._p.get_or_add_pPr()
                pBdr = OxmlElement("w:pBdr")
                bottom = OxmlElement("w:bottom")
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:sz"), "4")
                bottom.set(qn("w:color"), "CCCCCC")
                pBdr.append(bottom)
                pPr.append(pBdr)

            def _clear_cell(cell):
                for p in list(cell.paragraphs):
                    p._element.getparent().remove(p._element)

            # ── Full-width header ──────────────────────────────────────────────
            name = profile_data.get("full_name", "")
            if name:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _compact(p, before=0, after=1)
                run = p.add_run(name)
                run.font.name = "Georgia"
                run.font.size = Pt(18)
                run.font.bold = True
                run.font.color.rgb = NAVY

            headline = profile_data.get("headline", "")
            if headline:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _compact(p, before=0, after=1)
                run = p.add_run(headline)
                run.font.name = "Georgia"
                run.font.size = Pt(10)
                run.font.color.rgb = GRAY

            contact_parts = []
            if profile_data.get("email"):
                contact_parts.append(profile_data["email"])
            if profile_data.get("phone"):
                contact_parts.append(profile_data["phone"])
            loc = profile_data.get("location")
            if isinstance(loc, dict) and loc.get("city"):
                city = loc["city"]
                state = loc.get("state", "")
                contact_parts.append(f"{city}, {state}" if state else city)
            elif isinstance(loc, str) and loc:
                contact_parts.append(loc)
            if profile_data.get("linkedin_url"):
                contact_parts.append("LinkedIn")
            if profile_data.get("github_url"):
                contact_parts.append("GitHub")
            if profile_data.get("website_url"):
                contact_parts.append("Portfolio")

            if contact_parts:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _compact(p, before=0, after=4)
                run = p.add_run(" · ".join(contact_parts))
                run.font.name = "Georgia"
                run.font.size = Pt(9)
                run.font.color.rgb = MID
                pPr = p._p.get_or_add_pPr()
                pBdr = OxmlElement("w:pBdr")
                bottom = OxmlElement("w:bottom")
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:sz"), "8")
                bottom.set(qn("w:color"), "2C3E50")
                pBdr.append(bottom)
                pPr.append(pBdr)

            # ── Two-column table ───────────────────────────────────────────────
            table = doc.add_table(rows=1, cols=2)
            table.autofit = False

            # Letter (21.59cm) minus 1.016cm×2 margins = 19.56cm usable; 72/28 split
            # table.columns[i].width only updates <w:tblGrid>, not the per-cell
            # <w:tcW> that Word actually uses for a fixed-layout table — set both,
            # or Word renders an even 50/50 split regardless of these values.
            MAIN_COL_WIDTH = Cm(14.08)
            SIDE_COL_WIDTH = Cm(5.48)
            table.columns[0].width = MAIN_COL_WIDTH
            table.columns[1].width = SIDE_COL_WIDTH
            table.rows[0].cells[0].width = MAIN_COL_WIDTH
            table.rows[0].cells[1].width = SIDE_COL_WIDTH

            # Remove outer borders; add light gray inner vertical divider
            tbl = table._tbl
            tbl_pr = tbl.find(qn("w:tblPr"))
            if tbl_pr is None:
                tbl_pr = OxmlElement("w:tblPr")
                tbl.insert(0, tbl_pr)
            tbl_borders = OxmlElement("w:tblBorders")
            for side in ("top", "left", "bottom", "right", "insideH"):
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"), "none")
                tbl_borders.append(b)
            insideV = OxmlElement("w:insideV")
            insideV.set(qn("w:val"), "single")
            insideV.set(qn("w:sz"), "4")
            insideV.set(qn("w:color"), "DDDDDD")
            tbl_borders.append(insideV)
            tbl_pr.append(tbl_borders)

            # Cell vertical alignment: top
            for i in (0, 1):
                tc_pr = table.cell(0, i)._tc.get_or_add_tcPr()
                v_align = OxmlElement("w:vAlign")
                v_align.set(qn("w:val"), "top")
                tc_pr.append(v_align)

            # Add left padding to sidebar cell
            side_tc_pr = table.cell(0, 1)._tc.get_or_add_tcPr()
            tc_mar = OxmlElement("w:tcMar")
            left_mar = OxmlElement("w:left")
            left_mar.set(qn("w:w"), "200")
            left_mar.set(qn("w:type"), "dxa")
            tc_mar.append(left_mar)
            side_tc_pr.append(tc_mar)

            main = table.cell(0, 0)
            side = table.cell(0, 1)
            _clear_cell(main)
            _clear_cell(side)

            # ── Main column ───────────────────────────────────────────────────
            summary = profile_data.get("professional_summary", "")
            if summary:
                _section_title(main, "Professional Summary")
                p = main.add_paragraph()
                _compact(p, before=0, after=2)
                run = p.add_run(summary)
                run.font.name = "Georgia"
                run.font.size = Pt(8.5)
                run.font.color.rgb = DARK

            experiences = profile_data.get("work_experiences") or []
            if experiences:
                _section_title(main, "Experience")
                for exp in experiences:
                    p = main.add_paragraph()
                    _compact(p, before=3, after=0)
                    tr = p.add_run(exp.get("job_title", ""))
                    tr.font.name = "Georgia"
                    tr.font.size = Pt(9)
                    tr.font.bold = True
                    tr.font.color.rgb = DARK
                    start = exp.get("start_date", "")
                    end = "Present" if exp.get("is_current") else exp.get("end_date", "")
                    if start or end:
                        date_str = f"   {start} – {end}" if (start and end) else f"   {start or end}"
                        dr = p.add_run(date_str)
                        dr.font.name = "Georgia"
                        dr.font.size = Pt(8)
                        dr.font.color.rgb = GRAY
                        dr.font.italic = True

                    company = exp.get("company_name", "")
                    loc_exp = exp.get("location", {})
                    company_str = company
                    if isinstance(loc_exp, dict) and loc_exp.get("city"):
                        company_str += f" · {loc_exp['city']}"
                    if company_str:
                        cp = main.add_paragraph()
                        _compact(cp, before=0, after=0)
                        cr = cp.add_run(company_str)
                        cr.font.name = "Georgia"
                        cr.font.size = Pt(8.5)
                        cr.font.color.rgb = MID
                        cr.font.italic = True

                    desc = exp.get("description", "")
                    bullets = []
                    if exp.get("achievements"):
                        bullets = exp["achievements"]
                    elif desc:
                        lines = [l.strip() for l in desc.splitlines() if l.strip()]
                        if len(lines) > 1:
                            bullets = lines
                        else:
                            bullets = [pt.strip() for pt in re.split(r"(?<=\.)\s+", desc.strip()) if pt.strip()]
                    for bullet in bullets:
                        bp = main.add_paragraph()
                        _compact(bp, before=0, after=0)
                        bp.paragraph_format.left_indent = Pt(10)
                        bp.paragraph_format.first_line_indent = Pt(-10)
                        br = bp.add_run("• " + bullet)
                        br.font.name = "Georgia"
                        br.font.size = Pt(8)
                        br.font.color.rgb = DARK

            projects = profile_data.get("projects") or []
            if projects:
                _section_title(main, "Projects")
                for proj in projects:
                    p = main.add_paragraph()
                    p.paragraph_format.space_before = Pt(4)
                    pr = p.add_run(proj.get("title", ""))
                    pr.font.name = "Georgia"
                    pr.font.size = Pt(9.5)
                    pr.font.bold = True
                    pr.font.color.rgb = DARK
                    if proj.get("description"):
                        dp = main.add_paragraph()
                        dp.paragraph_format.space_before = Pt(0)
                        dp.paragraph_format.space_after = Pt(2)
                        dr = dp.add_run(proj["description"])
                        dr.font.name = "Georgia"
                        dr.font.size = Pt(8.5)
                        dr.font.color.rgb = MID

            # ── Sidebar column ────────────────────────────────────────────────
            educations = profile_data.get("educations") or []
            if educations:
                _section_title(side, "Education")
                for edu in educations:
                    p = side.add_paragraph()
                    _compact(p, before=3, after=0)
                    dr = p.add_run(edu.get("degree", ""))
                    dr.font.name = "Georgia"
                    dr.font.size = Pt(9)
                    dr.font.bold = True
                    dr.font.color.rgb = DARK
                    if edu.get("field_of_study"):
                        fp = side.add_paragraph()
                        _compact(fp, before=0, after=0)
                        fr = fp.add_run(edu["field_of_study"])
                        fr.font.name = "Georgia"
                        fr.font.size = Pt(8)
                        fr.font.color.rgb = MID
                    if edu.get("institution"):
                        ip = side.add_paragraph()
                        _compact(ip, before=0, after=0)
                        ir = ip.add_run(edu["institution"])
                        ir.font.name = "Georgia"
                        ir.font.size = Pt(8.5)
                        ir.font.color.rgb = MID
                    date_parts = []
                    if edu.get("start_date"):
                        date_parts.append(str(edu["start_date"])[:4])
                    if edu.get("end_date"):
                        date_parts.append(str(edu["end_date"])[:4])
                    if date_parts:
                        ep = side.add_paragraph()
                        _compact(ep, before=0, after=2)
                        er = ep.add_run(" – ".join(date_parts))
                        er.font.name = "Georgia"
                        er.font.size = Pt(8)
                        er.font.color.rgb = GRAY

            skills = profile_data.get("skills") or []
            if skills:
                _section_title(side, "Skills")
                from collections import defaultdict
                by_cat: dict[str, list[str]] = defaultdict(list)
                for s in skills:
                    by_cat[s.get("category", "general")].append(s.get("name", ""))
                for cat, names in by_cat.items():
                    gp = side.add_paragraph()
                    _compact(gp, before=3, after=0)
                    gr = gp.add_run(cat.replace("_", " ").title())
                    gr.font.name = "Georgia"
                    gr.font.size = Pt(8.5)
                    gr.font.bold = True
                    gr.font.color.rgb = NAVY
                    np_ = side.add_paragraph()
                    _compact(np_, before=0, after=1)
                    nr = np_.add_run(", ".join(names))
                    nr.font.name = "Georgia"
                    nr.font.size = Pt(8.5)
                    nr.font.color.rgb = DARK

            certs = profile_data.get("certifications") or []
            if certs:
                _section_title(side, "Certifications")
                for cert in certs:
                    cp = side.add_paragraph()
                    _compact(cp, before=3, after=0)
                    cr = cp.add_run(cert.get("name", ""))
                    cr.font.name = "Georgia"
                    cr.font.size = Pt(8.5)
                    cr.font.bold = True
                    cr.font.color.rgb = DARK
                    if cert.get("issuing_organization"):
                        op = side.add_paragraph()
                        _compact(op, before=0, after=1)
                        or_ = op.add_run(cert["issuing_organization"])
                        or_.font.name = "Georgia"
                        or_.font.size = Pt(8)
                        or_.font.color.rgb = GRAY

            achievements = profile_data.get("achievements") or []
            if achievements:
                _section_title(side, "Achievements")
                for ach in achievements:
                    ap = side.add_paragraph(style="List Bullet")
                    ar = ap.add_run(ach.get("title", ""))
                    ar.font.name = "Georgia"
                    ar.font.size = Pt(8.5)

            buffer = io.BytesIO()
            doc.save(buffer)
            return buffer.getvalue()

        # ── Dynamic Layout (Default/ATS/Modern/Minimal/Executive) ─────────────
        font_name = "Arial"
        heading_color = RGBColor(0x0F, 0x17, 0x2A)
        header_align = WD_ALIGN_PARAGRAPH.CENTER
        heading_align = WD_ALIGN_PARAGRAPH.LEFT

        if template_slug == "creative_minimal":
            font_name = "Georgia"
            heading_align = WD_ALIGN_PARAGRAPH.CENTER
            header_align = WD_ALIGN_PARAGRAPH.CENTER
            heading_color = RGBColor(0x99, 0x99, 0x99)
        elif template_slug == "executive_pro":
            font_name = "Times New Roman"
            heading_color = RGBColor(0x0F, 0x17, 0x2A)
            header_align = WD_ALIGN_PARAGRAPH.LEFT
        elif template_slug == "professional":
            font_name = "Georgia"
            heading_color = RGBColor(0x2C, 0x3E, 0x50)
            header_align = WD_ALIGN_PARAGRAPH.CENTER
        elif template_slug == "modern_tech":
            font_name = "Arial"
            heading_color = RGBColor(0x11, 0x18, 0x27)
            header_align = WD_ALIGN_PARAGRAPH.LEFT
        else:
            header_align = WD_ALIGN_PARAGRAPH.CENTER

        doc.styles["Normal"].font.name = font_name
        
        from docx.enum.text import WD_LINE_SPACING
        normal_p_fmt = doc.styles["Normal"].paragraph_format
        normal_p_fmt.space_before = Pt(0)
        normal_p_fmt.space_after = Pt(2)
        normal_p_fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
        normal_p_fmt.line_spacing = 1.0

        TEXT_WIDTH_TWIPS = 9936
        for section in doc.sections:
            section.top_margin = Cm(1.2)
            section.bottom_margin = Cm(1.2)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)

        def add_right_tab(p):
            pPr = p._p.get_or_add_pPr()
            tabs_elem = OxmlElement("w:tabs")
            tab_elem = OxmlElement("w:tab")
            tab_elem.set(qn("w:val"), "right")
            tab_elem.set(qn("w:pos"), str(TEXT_WIDTH_TWIPS))
            tabs_elem.append(tab_elem)
            pPr.append(tabs_elem)

        def add_hyperlink(para, url, text, font_size_pt=9):
            part = para.part
            r_id = part.relate_to(
                url,
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                is_external=True,
            )
            hl = OxmlElement("w:hyperlink")
            hl.set(qn("r:id"), r_id)
            r = OxmlElement("w:r")
            rPr = OxmlElement("w:rPr")
            color = OxmlElement("w:color")
            color.set(qn("w:val"), "0F172A")
            rPr.append(color)
            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), str(int(font_size_pt * 2)))
            rPr.append(sz)
            u = OxmlElement("w:u")
            u.set(qn("w:val"), "single")
            rPr.append(u)
            r.append(rPr)
            t = OxmlElement("w:t")
            t.text = text
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            r.append(t)
            hl.append(r)
            para._p.append(hl)

        def add_section_heading(title: str):
            p = doc.add_paragraph()
            p.alignment = heading_align
            run = p.add_run(title.upper())
            run.font.name = font_name
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = heading_color
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:color"), "DDDDDD" if template_slug == "creative_minimal" else "000000")
            pBdr.append(bottom)
            pPr.append(pBdr)

        # ── Header: Name & headline ──────────────────────────────────────────
        name = profile_data.get("full_name", "")
        if name:
            p = doc.add_paragraph()
            p.alignment = header_align
            run = p.add_run(name)
            run.font.name = font_name
            run.font.size = Pt(20)
            run.font.bold = True
            run.font.color.rgb = heading_color

        headline = profile_data.get("headline", "")
        if headline:
            p = doc.add_paragraph()
            p.alignment = header_align
            run = p.add_run(headline)
            run.font.name = font_name
            run.font.size = Pt(11)
            run.font.italic = True

        # Contact: text info left, URL links right-aligned
        text_parts: list[str] = []
        if profile_data.get("email"):
            text_parts.append(profile_data["email"])
        if profile_data.get("phone"):
            text_parts.append(profile_data["phone"])
        loc = profile_data.get("location")
        if isinstance(loc, dict) and loc.get("city"):
            city = loc["city"]
            state = loc.get("state", "")
            text_parts.append(f"{city}, {state}" if state else city)
        elif isinstance(loc, str) and loc:
            text_parts.append(loc)

        link_items: list[tuple[str, str]] = []
        if profile_data.get("linkedin_url"):
            link_items.append(("LinkedIn", profile_data["linkedin_url"]))
        if profile_data.get("github_url"):
            link_items.append(("GitHub", profile_data["github_url"]))
        if profile_data.get("twitter_url"):
            link_items.append(("X / Twitter", profile_data["twitter_url"]))
        if profile_data.get("website_url"):
            link_items.append(("Portfolio", profile_data["website_url"]))

        if text_parts or link_items:
            p = doc.add_paragraph()
            p.alignment = header_align if not link_items else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(4)
            if link_items and header_align != WD_ALIGN_PARAGRAPH.CENTER:
                add_right_tab(p)
            
            if header_align == WD_ALIGN_PARAGRAPH.CENTER:
                all_parts = text_parts + [l[0] for l in link_items]
                lr = p.add_run("  |  ".join(all_parts))
                lr.font.name = font_name
                lr.font.size = Pt(9)
                lr.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
            else:
                if text_parts:
                    lr = p.add_run("  |  ".join(text_parts))
                    lr.font.name = font_name
                    lr.font.size = Pt(9)
                    lr.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
                if link_items:
                    r = p.add_run("\t")
                    r.font.name = font_name
                    r.font.size = Pt(9)
                    for i, (label, url) in enumerate(link_items):
                        if i > 0:
                            sep = p.add_run("  ·  ")
                            sep.font.name = font_name
                            sep.font.size = Pt(9)
                            sep.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
                        add_hyperlink(p, url, label)

        # ── Summary ──────────────────────────────────────────────────────────
        summary = profile_data.get("professional_summary", "")
        if summary:
            add_section_heading("Professional Summary")
            p = doc.add_paragraph(summary)
            if p.runs:
                p.runs[0].font.name = font_name
            p.paragraph_format.space_after = Pt(4)

        # ── Work Experience ───────────────────────────────────────────────────
        experiences = profile_data.get("work_experiences") or []
        if experiences:
            add_section_heading("Work Experience")
            for exp in experiences:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(0)
                add_right_tab(p)
                title_run = p.add_run(exp.get("job_title", ""))
                title_run.font.name = font_name
                title_run.font.bold = True
                title_run.font.size = Pt(10.5)
                start = exp.get("start_date", "")
                end = "Present" if exp.get("is_current") else exp.get("end_date", "")
                if start or end:
                    dr = p.add_run(f"\t{start} – {end}")
                    dr.font.name = font_name
                    dr.font.size = Pt(9)
                    dr.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
                    dr.font.italic = True

                company = exp.get("company_name", "")
                location = exp.get("location", {})
                company_str = company
                if isinstance(location, dict) and location.get("city"):
                    company_str += f"  |  {location['city']}"
                if company_str:
                    cp = doc.add_paragraph()
                    cp.paragraph_format.space_before = Pt(0)
                    cp.paragraph_format.space_after = Pt(2)
                    cr = cp.add_run(company_str)
                    cr.font.name = font_name
                    cr.font.size = Pt(10)
                    cr.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

                desc = exp.get("description", "")
                bullet_texts = []
                if desc:
                    lines = [l.strip() for l in desc.splitlines() if l.strip()]
                    if len(lines) > 1:
                        bullet_texts = lines
                    else:
                        parts = re.split(r"(?<=\.)\s+", desc.strip())
                        bullet_texts = [pt.strip() for pt in parts if pt.strip()]

                for bullet in bullet_texts:
                    bp = doc.add_paragraph(style="List Bullet")
                    r = bp.add_run(bullet)
                    r.font.name = font_name
                    r.font.size = Pt(9.5)
                    bp.paragraph_format.space_before = Pt(0)
                    bp.paragraph_format.space_after = Pt(1)

                for ach in (exp.get("achievements") or []):
                    ap = doc.add_paragraph(style="List Bullet")
                    r = ap.add_run(ach)
                    r.font.name = font_name
                    r.font.size = Pt(9.5)
                    ap.paragraph_format.space_before = Pt(0)
                    ap.paragraph_format.space_after = Pt(1)

        # ── Education ────────────────────────────────────────────────────────
        educations = profile_data.get("educations") or []
        if educations:
            add_section_heading("Education")
            for edu in educations:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(0)
                add_right_tab(p)
                deg = edu.get("degree", "")
                field = edu.get("field_of_study", "")
                degree_str = f"{deg} in {field}" if field else deg
                dr = p.add_run(degree_str)
                dr.font.name = font_name
                dr.font.bold = True
                dr.font.size = Pt(10.5)
                start_yr = str(edu.get("start_date", "") or "")[:4]
                end_yr = str(edu.get("end_date", "") or "")[:4]
                date_str = (
                    f"{start_yr} – {end_yr}" if (start_yr and end_yr)
                    else (start_yr or end_yr)
                )
                if date_str:
                    er = p.add_run(f"\t{date_str}")
                    er.font.name = font_name
                    er.font.size = Pt(9)
                    er.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
                    er.font.italic = True

                inst = edu.get("institution", "")
                if inst:
                    ip = doc.add_paragraph()
                    ip.paragraph_format.space_before = Pt(0)
                    ip.paragraph_format.space_after = Pt(2)
                    ir = ip.add_run(inst)
                    ir.font.name = font_name
                    ir.font.size = Pt(10)
                    ir.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

                if edu.get("gpa"):
                    gp = doc.add_paragraph(f"GPA: {edu['gpa']}")
                    if gp.runs:
                        gp.runs[0].font.name = font_name
                    gp.paragraph_format.space_after = Pt(1)

        # ── Skills ───────────────────────────────────────────────────────────
        skills = profile_data.get("skills") or []
        if skills:
            add_section_heading("Skills")
            by_cat = {}
            for s in skills:
                cat = s.get("category", "general").replace("_", " ").title()
                if cat not in by_cat:
                    by_cat[cat] = []
                by_cat[cat].append(s.get("name", ""))

            for cat, items in by_cat.items():
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                if template_slug == "modern_tech":
                    # Hanging indent for left-aligned categories
                    p.paragraph_format.left_indent = Cm(3.5)
                    p.paragraph_format.first_line_indent = -Cm(3.5)
                    p.paragraph_format.tab_stops.add_tab_stop(Cm(3.5))
                    cr = p.add_run(f"{cat}\t")
                    cr.font.name = "Courier New"
                    cr.font.size = Pt(9)
                    cr.font.bold = True
                    sr = p.add_run(", ".join(items))
                    sr.font.name = "Courier New"
                    sr.font.size = Pt(9)
                    sr.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
                elif template_slug == "creative_minimal":
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cr = p.add_run(f"{cat}: ")
                    cr.font.name = font_name
                    cr.font.bold = True
                    cr.font.size = Pt(9)
                    sr = p.add_run(", ".join(items))
                    sr.font.name = font_name
                    sr.font.size = Pt(9)
                else:
                    # ATS Clean / Professional layouts: add hanging indent
                    p.paragraph_format.left_indent = Cm(3.0)
                    p.paragraph_format.first_line_indent = -Cm(3.0)
                    p.paragraph_format.tab_stops.add_tab_stop(Cm(3.0))
                    cr = p.add_run(f"{cat}:\t")
                    cr.font.name = font_name
                    cr.font.bold = True
                    cr.font.size = Pt(9.5)
                    sr = p.add_run(", ".join(items))
                    sr.font.name = font_name
                    sr.font.size = Pt(9.5)

        # ── Projects ─────────────────────────────────────────────────────────
        projects = profile_data.get("projects") or []
        if projects:
            add_section_heading("Projects")
            for proj in projects:
                p = doc.add_paragraph()
                r = p.add_run(proj.get("title", ""))
                r.font.name = font_name
                r.font.bold = True
                if proj.get("description"):
                    dp = doc.add_paragraph(proj["description"])
                    if dp.runs:
                        dp.runs[0].font.name = font_name
                for hl in (proj.get("highlights") or []):
                    p = doc.add_paragraph(style="List Bullet")
                    r = p.add_run(hl)
                    r.font.name = font_name
                    r.font.size = Pt(10)

        # ── Certifications ────────────────────────────────────────────────────
        certs = profile_data.get("certifications") or []
        if certs:
            add_section_heading("Certifications")
            for cert in certs:
                p = doc.add_paragraph()
                r = p.add_run(cert.get("name", ""))
                r.font.name = font_name
                r.font.bold = True
                org = cert.get("issuing_organization", "")
                if org:
                    r2 = p.add_run(f"  —  {org}")
                    r2.font.name = font_name

        # ── Achievements ──────────────────────────────────────────────────────
        achievements = profile_data.get("achievements") or []
        if achievements:
            add_section_heading("Achievements")
            for ach in achievements:
                p = doc.add_paragraph(style="List Bullet")
                r = p.add_run(ach.get("title", ""))
                r.font.name = font_name

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
