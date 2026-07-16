import logging
import uuid
import re
from datetime import timedelta
from django.db import transaction
from django.utils import timezone
from apps.exports.models import ExportJob, ExportStatus, ExportResourceType, ExportFormat

logger = logging.getLogger(__name__)

# Download URL validity: 24 hours
DOWNLOAD_URL_TTL_SECONDS = 86400

CONTENT_TYPES = {
    ExportFormat.PDF: "application/pdf",
    ExportFormat.DOCX: "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}

FILE_EXTENSIONS = {
    ExportFormat.PDF: "pdf",
    ExportFormat.DOCX: "docx",
}


class ExportService:

    # ── Request ───────────────────────────────────────────────────────────────

    @staticmethod
    @transaction.atomic
    def request(user, resource_type: str, resource_id: str, fmt: str, template_slug: str = None) -> ExportJob:
        """Create an ExportJob and enqueue async generation."""
        template = None
        if template_slug:
            try:
                from apps.templates_engine.models import Template
                template = Template.objects.get(slug=template_slug, is_active=True)
            except Exception:
                pass

        job = ExportJob.objects.create(
            user=user,
            resource_type=resource_type,
            resource_id=resource_id,
            format=fmt,
            template=template,
            status=ExportStatus.QUEUED,
        )

        try:
            from celery_app.tasks.export_tasks import generate_export
            generate_export.delay(str(job.id))
        except Exception:
            logger.warning("Celery unavailable — generating export synchronously")
            ExportService.generate(job)

        return job

    # ── Generate (called by Celery task) ─────────────────────────────────────

    @staticmethod
    def generate(job: ExportJob) -> None:
        """Render the resource and upload the generated file to storage."""
        from apps.exports.generators.registry import PdfGenerator, DocxGenerator  # ensure registration
        from apps.exports.generators import get_generator

        job.status = ExportStatus.PROCESSING
        job.save(update_fields=["status"])

        try:
            profile_data, html_content = ExportService._render_resource(job)

            if job.resource_type == ExportResourceType.COVER_LETTER and job.format == ExportFormat.DOCX:
                file_bytes = ExportService._generate_cover_letter_docx(job, profile_data)
            else:
                generator = get_generator(job.format)
                template_slug = job.template.slug if job.template else None
                file_bytes = generator.generate(html_content, profile_data=profile_data, template_slug=template_slug)

            ext = FILE_EXTENSIONS[job.format]
            storage_path = f"exports/{job.user_id}/{job.id}.{ext}"
            content_type = CONTENT_TYPES[job.format]

            from storage import storage as get_storage
            store = get_storage()
            store.upload(storage_path, file_bytes, content_type=content_type)

            expires_at = timezone.now() + timedelta(seconds=DOWNLOAD_URL_TTL_SECONDS)
            download_url = store.get_signed_url(storage_path, expiry_seconds=DOWNLOAD_URL_TTL_SECONDS)

            job.file_path = storage_path
            job.file_size = len(file_bytes)
            job.download_url = download_url
            job.url_expires_at = expires_at
            job.status = ExportStatus.COMPLETED
            job.completed_at = timezone.now()
            job.save(update_fields=[
                "file_path", "file_size", "download_url", "url_expires_at",
                "status", "completed_at",
            ])

        except Exception as exc:
            job.status = ExportStatus.FAILED
            job.error_message = str(exc)[:1000]
            job.save(update_fields=["status", "error_message"])
            raise

    @staticmethod
    def _render_resource(job: ExportJob) -> tuple[dict, str]:
        """Fetch resource + profile, render HTML. Returns (profile_data, html)."""
        from apps.profiles.profile_utils import ProfileSerializer
        from apps.profiles.models import UserProfile

        profile = UserProfile.objects.prefetch_related(
            "work_experiences", "educations", "skills",
            "projects", "certifications", "achievements", "publications",
        ).get(user=job.user)
        profile_data = ProfileSerializer.to_dict(profile)

        html = ""
        if job.resource_type == ExportResourceType.RESUME:
            html = ExportService._render_resume(job, profile_data)
        elif job.resource_type == ExportResourceType.COVER_LETTER:
            html = ExportService._render_cover_letter(job, profile_data)
        elif job.resource_type == ExportResourceType.PORTFOLIO:
            html = ExportService._render_portfolio(job, profile_data)

        return profile_data, html

    @staticmethod
    def _render_resume(job: ExportJob, profile_data: dict) -> str:
        from apps.resumes.models import Resume
        from apps.templates_engine.services import TemplateRenderer

        resume = Resume.objects.select_related("template").get(
            id=job.resource_id, user=job.user
        )
        template = job.template or resume.template
        if not template:
            from apps.templates_engine.models import Template, TemplateType
            template = Template.objects.filter(type=TemplateType.RESUME, is_active=True).first()
        if not template:
            return "<p>No template available.</p>"

        return TemplateRenderer.render_resume(template, profile_data, resume.template_settings)

    @staticmethod
    def _render_cover_letter(job: ExportJob, profile_data: dict) -> str:
        from apps.cover_letters.models import CoverLetter
        from apps.cover_letters.services import CoverLetterService

        letter = CoverLetter.objects.select_related("template").get(
            id=job.resource_id, user=job.user
        )
        template = job.template or letter.template
        if not template:
            from apps.templates_engine.models import Template, TemplateType
            template = Template.objects.filter(type=TemplateType.COVER_LETTER, is_active=True).first()
        if not template:
            return "<p>No template available.</p>"

        return CoverLetterService._render_html_with_template(letter, profile_data, template)

    @staticmethod
    def _generate_cover_letter_docx(job: ExportJob, profile_data: dict) -> bytes:
        from apps.cover_letters.models import CoverLetter

        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            from docx.shared import Cm, Pt, RGBColor
            import io
        except ImportError:
            raise RuntimeError(
                "python-docx is required for DOCX export. "
                "Install it: pip install python-docx"
            )

        letter = CoverLetter.objects.select_related("template").get(
            id=job.resource_id, user=job.user
        )
        doc = Document()
        template_slug = letter.template.slug if letter.template else ""

        # Default to Modern ATS styles
        t_marg, b_marg, l_marg, r_marg = 1.9, 1.9, 2.54, 2.54
        font_name = "Arial"
        base_size = 11

        if template_slug == "professional-letter" or template_slug == "professional_letter":
            font_name = "Georgia"
            t_marg, b_marg = 2.03, 2.03
        elif template_slug == "executive-letter" or template_slug == "executive_letter":
            t_marg, b_marg = 1.9, 1.9
        elif template_slug == "technical-letter" or template_slug == "technical_letter":
            t_marg, b_marg = 1.65, 1.65

        for section in doc.sections:
            section.top_margin = Cm(t_marg)
            section.bottom_margin = Cm(b_marg)
            section.left_margin = Cm(l_marg)
            section.right_margin = Cm(r_marg)

        def apply_font(run, size=None, bold=False, color=0x00_00_00):
            run.font.name = font_name
            run.font.size = Pt(size or base_size)
            run.font.bold = bold
            run.font.color.rgb = RGBColor(
                (color >> 16) & 0xFF,
                (color >> 8) & 0xFF,
                color & 0xFF,
            )

        def compact(paragraph, before=0, after=6, line_spacing=1.12):
            paragraph.paragraph_format.space_before = Pt(before)
            paragraph.paragraph_format.space_after = Pt(after)
            paragraph.paragraph_format.line_spacing = line_spacing

        def add_rule(paragraph):
            p_pr = paragraph._p.get_or_add_pPr()
            p_bdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "000000")
            p_bdr.append(bottom)
            p_pr.append(p_bdr)

        name = profile_data.get("full_name") or job.user.get_full_name() or job.user.email
        p = doc.add_paragraph()
        compact(p, after=4)
        run = p.add_run(name)
        apply_font(run, size=18, bold=True)

        contact_items = []
        if profile_data.get("email"):
            contact_items.append(("✉ ", profile_data["email"]))
        if profile_data.get("phone"):
            contact_items.append(("📞 ", profile_data["phone"]))
        location = profile_data.get("location")
        if isinstance(location, dict):
            city = location.get("city")
            state = location.get("state")
            if city:
                contact_items.append(("📍 ", f"{city}, {state}" if state else city))
        elif location:
            contact_items.append(("📍 ", str(location)))
            
        if profile_data.get("linkedin_url"):
            linkedin = profile_data["linkedin_url"].replace("https://", "").replace("www.", "")
            contact_items.append(("🔗 ", linkedin))

        if contact_items:
            p = doc.add_paragraph()
            compact(p, after=8)
            for i, (icon, text) in enumerate(contact_items):
                if i > 0:
                    run_sep = p.add_run("  |  ")
                    apply_font(run_sep, size=9, color=0x94_A3_B8)
                
                run_icon = p.add_run(icon)
                apply_font(run_icon, size=9, color=0x1D_4E_D8)
                
                run_text = p.add_run(text)
                apply_font(run_text, size=9, color=0x33_33_33)
            
            # Apply divider if modern or executive
            if "modern" in template_slug or "executive" in template_slug or not template_slug:
                add_rule(p)

        p = doc.add_paragraph()
        compact(p, before=10, after=12)
        run = p.add_run(timezone.now().strftime("%B %d, %Y"))
        apply_font(run, size=None)

        recipient_lines = [
            letter.hiring_manager_name,
            letter.hiring_manager_title,
            letter.company_name,
        ]
        company_address = letter.company_address or {}
        if isinstance(company_address, dict):
            city = company_address.get("city")
            state = company_address.get("state")
            if city:
                recipient_lines.append(f"{city}, {state}" if state else city)

        for index, line in enumerate([line for line in recipient_lines if line]):
            p = doc.add_paragraph()
            compact(p, after=0 if index < 2 else 10)
            run = p.add_run(line)
            apply_font(run, size=None, bold=index in (0, 2))

        salutation = (
            f"Dear {letter.hiring_manager_name},"
            if letter.hiring_manager_name
            else "Dear Hiring Manager,"
        )
        p = doc.add_paragraph()
        compact(p, before=8, after=14)
        run = p.add_run(salutation)
        apply_font(run, size=None)

        paragraphs = [part.strip() for part in (letter.body_content or "").split("\n\n") if part.strip()]
        for paragraph in paragraphs:
            p = doc.add_paragraph()
            compact(p, after=12, line_spacing=1.18)
            run = p.add_run(paragraph)
            apply_font(run, size=None)

        p = doc.add_paragraph()
        compact(p, before=4, after=24)
        run = p.add_run("Sincerely,")
        apply_font(run, size=None)

        p = doc.add_paragraph()
        compact(p, after=0)
        run = p.add_run(name)
        apply_font(run, size=None, bold=True)

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    @staticmethod
    def _render_portfolio(job: ExportJob, profile_data: dict) -> str:
        from apps.portfolios.models import Portfolio
        from apps.portfolios.services import PortfolioService

        portfolio = Portfolio.objects.get(user=job.user)
        return PortfolioService.rebuild_preview(portfolio)

    # ── Download URL ──────────────────────────────────────────────────────────

    @staticmethod
    def get_download_url(job: ExportJob) -> str:
        """Return a valid signed download URL. Regenerates if expired."""
        if job.status != ExportStatus.COMPLETED:
            raise ValueError("Export is not ready for download.")

        now = timezone.now()
        if job.url_expires_at and job.url_expires_at > now:
            return job.download_url

        # Regenerate
        from storage import storage as get_storage
        store = get_storage()
        expires_at = now + timedelta(seconds=DOWNLOAD_URL_TTL_SECONDS)
        new_url = store.get_signed_url(job.file_path, expiry_seconds=DOWNLOAD_URL_TTL_SECONDS)
        job.download_url = new_url
        job.url_expires_at = expires_at
        job.save(update_fields=["download_url", "url_expires_at"])
        return new_url

    @staticmethod
    def get_resource_title(job: ExportJob) -> str:
        try:
            if job.resource_type == ExportResourceType.COVER_LETTER:
                from apps.cover_letters.models import CoverLetter
                return CoverLetter.objects.get(id=job.resource_id, user=job.user).title
            if job.resource_type == ExportResourceType.RESUME:
                from apps.resumes.models import Resume
                return Resume.objects.get(id=job.resource_id, user=job.user).title
            if job.resource_type == ExportResourceType.PORTFOLIO:
                return "Portfolio"
        except Exception:
            pass
        return job.resource_type.replace("_", " ").title()

    @staticmethod
    def get_download_filename(job: ExportJob) -> str:
        ext = FILE_EXTENSIONS.get(job.format, job.format)
        title = ExportService.get_resource_title(job)
        stem = re.sub(r"[^A-Za-z0-9._ -]+", "", title).strip()
        stem = re.sub(r"\s+", " ", stem).strip(" ._-")
        if not stem:
            stem = job.resource_type.replace("_", "-")
        return f"{stem}.{ext}"

    # ── Regenerate ────────────────────────────────────────────────────────────

    @staticmethod
    def regenerate(job: ExportJob) -> ExportJob:
        """Create a new export job for the same resource (re-export with latest data)."""
        return ExportService.request(
            user=job.user,
            resource_type=job.resource_type,
            resource_id=str(job.resource_id),
            fmt=job.format,
            template_slug=job.template.slug if job.template else None,
        )

    # ── Query ─────────────────────────────────────────────────────────────────

    @staticmethod
    def list_for_user(user, resource_type: str | None = None) -> list:
        qs = ExportJob.objects.filter(user=user)
        if resource_type:
            qs = qs.filter(resource_type=resource_type)
        return qs.select_related("template").order_by("-created_at")

    @staticmethod
    def get_for_user(job_id: str, user) -> ExportJob:
        return ExportJob.objects.select_related("template").get(id=job_id, user=user)

    @staticmethod
    def delete(job: ExportJob) -> None:
        if job.file_path:
            try:
                from storage import storage as get_storage
                get_storage().delete(job.file_path)
            except Exception:
                pass
        job.delete()
